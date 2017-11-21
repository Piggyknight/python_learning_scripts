from docx import Document
from  docx.shared import  Pt

def Export(filePath, contentLines, fontName, fontSize):
    '''
    :param contentLines: whole content for the word doc
    :return: no return
    '''
    document = Document()
    p = document.add_paragraph()
    for line in contentLines:
        run = p.add_run(line)
        run.font.size = Pt(fontSize)
        run.font.name = fontName

    document.save(filePath)
    return



